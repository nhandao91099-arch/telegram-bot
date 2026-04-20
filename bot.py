from telegram import Update, InlineKeyboardMarkup, InlineKeyboardButton
from telegram.ext import ApplicationBuilder, CommandHandler, CallbackQueryHandler, MessageHandler, filters, ContextTypes

from docx import Document
from reportlab.platypus import SimpleDocTemplate, Image
from PIL import Image as PILImage
import pytesseract
import zipfile
import os
import re
from datetime import datetime

# OCR
pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"

TOKEN = "DAN_TOKEN_VAO_DAY"

user_data = {}
REQUIRED_DOCS = ["cccd","vneid","luong","vssid","sf","other"]

# ================= START =================
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    kb = [
        [InlineKeyboardButton("📊 TÍNH LÃI", callback_data="tls")],
        [InlineKeyboardButton("📄 THƯ XÁC NHẬN", callback_data="form")],
        [InlineKeyboardButton("📂 GIẤY TỜ", callback_data="docs")]
    ]
    await update.message.reply_text("Chọn:", reply_markup=InlineKeyboardMarkup(kb))

# ================= TLS =================
def tinh_lai(luong):
    if luong <= 10: return 45
    elif luong <= 20: return 40
    return 35

# ================= DBR =================
def tinh_dbr(luong, kv, no):
    max_dbr = 0.5 if kv=="tp" else 0.6
    con = luong*max_dbr - no
    vay = int(con*15)
    return round(no/luong*100), int(vay)

# ================= OCR =================
def ocr_image(path):
    return pytesseract.image_to_string(PILImage.open(path))

# ================= PDF =================
def tao_pdf(docs, context, uid):
    files=[]
    for loai,fid in docs.items():
        file=context.bot.get_file(fid)
        img=f"{uid}_{loai}.jpg"
        file.download(img)

        pdf=f"{uid}_{loai}.pdf"
        SimpleDocTemplate(pdf).build([Image(img,400,500)])
        files.append(pdf)

        try: os.remove(img)
        except: pass

    zipf=f"{uid}.zip"
    with zipfile.ZipFile(zipf,'w') as z:
        for f in files: z.write(f)

    for f in files:
        try: os.remove(f)
        except: pass

    return zipf

# ================= FORM =================
def tao_form(data, uid):
    doc=Document()
    now=datetime.now()

    doc.add_paragraph(f"Tên: {data.get('ten','')}")
    doc.add_paragraph(f"CCCD: {data.get('cccd','')}")
    doc.add_paragraph(f"Ngày cấp: {data.get('ngay','')}")
    doc.add_paragraph(f"Nơi cấp: {data.get('noicap','')}")
    doc.add_paragraph(f"Địa chỉ: {data.get('dc','')}")

    doc.add_paragraph(f"\nNgày {now.day} tháng {now.month} năm {now.year}")

    path=f"{uid}_form.docx"
    doc.save(path)
    return path

# ================= BUTTON =================
async def button(update: Update, context: ContextTypes.DEFAULT_TYPE):
    q=update.callback_query
    uid=q.from_user.id
    data=q.data

    user_data.setdefault(uid,{})
    await q.answer()

    if data=="tls":
        user_data[uid]["step"]="luong"
        await q.message.reply_text("Nhập lương:")
        return

    if data=="form":
        user_data[uid]["step"]="form"
        await q.message.reply_text("Nhập:\nTên\nCCCD\nNgày cấp\nNơi cấp\nĐịa chỉ")
        return

    if data=="docs":
        kb=[[InlineKeyboardButton(x.upper(),callback_data=f"doc_{x}")] for x in REQUIRED_DOCS]
        kb.append([InlineKeyboardButton("CHECK",callback_data="check")])
        kb.append([InlineKeyboardButton("PDF",callback_data="pdf")])
        kb.append([InlineKeyboardButton("SCAN",callback_data="scan")])
        await q.message.reply_text("Chọn:",reply_markup=InlineKeyboardMarkup(kb))
        return

    if data.startswith("doc_"):
        user_data[uid]["doc"]=data.split("_")[1]
        user_data[uid]["step"]="upload"
        await q.message.reply_text("Gửi ảnh")
        return

    if data=="check":
        docs=user_data[uid].get("docs",{})
        txt="\n".join([f"{'✅' if d in docs else '❌'} {d}" for d in REQUIRED_DOCS])
        await q.message.reply_text(txt)
        return

    if data=="pdf":
        docs=user_data[uid].get("docs",{})
        zipf=tao_pdf(docs,context,uid)
        await q.message.reply_document(open(zipf,"rb"))
        try: os.remove(zipf)
        except: pass
        return

    if data=="scan":
        docs=user_data[uid].get("docs",{})
        text=""
        for loai,fid in docs.items():
            file=context.bot.get_file(fid)
            path=f"{uid}_{loai}.jpg"
            file.download(path)
            text+=ocr_image(path)
            try: os.remove(path)
            except: pass

        cccd=re.search(r"\d{9,12}",text)
        await q.message.reply_text(f"CCCD: {cccd.group() if cccd else 'không rõ'}")
        return

# ================= HANDLE TEXT =================
async def handle(update: Update, context: ContextTypes.DEFAULT_TYPE):
    uid=update.message.from_user.id
    txt=update.message.text
    step=user_data.get(uid,{}).get("step")

    if step=="luong":
        luong=float(txt)
        user_data[uid]["luong"]=luong
        user_data[uid]["step"]="no"
        await update.message.reply_text("Nhập nợ:")
        return

    if step=="no":
        no=float(txt)
        luong=user_data[uid]["luong"]
        dbr,vay=tinh_dbr(luong,"tp",no)

        kb=[[InlineKeyboardButton("📂 GIẤY TỜ",callback_data="docs")]]
        await update.message.reply_text(f"DBR: {dbr}%\nVay: {vay}tr",reply_markup=InlineKeyboardMarkup(kb))
        return

    if step=="form":
        lines=txt.split("\n")
        if len(lines)<5:
            await update.message.reply_text("Nhập đủ 5 dòng")
            return

        data={
            "ten":lines[0],
            "cccd":lines[1],
            "ngay":lines[2],
            "noicap":lines[3],
            "dc":lines[4]
        }

        path=tao_form(data,uid)
        await update.message.reply_document(open(path,"rb"))

        try: os.remove(path)
        except: pass

# ================= PHOTO =================
async def photo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    uid=update.message.from_user.id

    if user_data.get(uid,{}).get("step")=="upload":
        fid=update.message.photo[-1].file_id
        loai=user_data[uid]["doc"]

        user_data[uid].setdefault("docs",{})
        user_data[uid]["docs"][loai]=fid

        await update.message.reply_text(f"Đã lưu {loai}")

# ================= RUN =================
app=ApplicationBuilder().token(TOKEN).build()

app.add_handler(CommandHandler("start",start))
app.add_handler(CallbackQueryHandler(button))
app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND,handle))
app.add_handler(MessageHandler(filters.PHOTO,photo))

app.run_polling()
